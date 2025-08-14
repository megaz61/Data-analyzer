import pandas as pd
import numpy as np  # Added missing import
import PyPDF2
import docx
import json
import os
from typing import Dict, Any, List
import uuid

class FileProcessor:
    def __init__(self):
        self.supported_formats = ['.csv', '.xlsx', '.xls', '.pdf', '.docx', '.txt']
    
    def process_file(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process uploaded file and extract data"""
        file_extension = os.path.splitext(original_filename)[1].lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        file_id = str(uuid.uuid4())
        
        try:
            if file_extension in ['.csv']:
                return self._process_csv(file_path, file_id, original_filename)
            elif file_extension in ['.xlsx', '.xls']:
                return self._process_excel(file_path, file_id, original_filename)
            elif file_extension == '.pdf':
                return self._process_pdf(file_path, file_id, original_filename)
            elif file_extension == '.docx':
                return self._process_docx(file_path, file_id, original_filename)
            elif file_extension == '.txt':
                return self._process_txt(file_path, file_id, original_filename)
        except Exception as e:
            raise Exception(f"Error processing file: {str(e)}")
    
    def _process_csv(self, file_path: str, file_id: str, filename: str) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            df = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    used_encoding = encoding
                    break
                except Exception as e:
                    print(f"Failed with encoding {encoding}: {e}")
                    continue
            
            if df is None:
                raise Exception("Could not read CSV file with any supported encoding")
            
            print(f"CSV loaded successfully with encoding: {used_encoding}")
            
            # Basic analysis
            analysis = self._analyze_dataframe(df)
            
            return {
                'file_id': file_id,
                'filename': filename,
                'type': 'csv',
                'data': df.to_dict('records')[:100],  # Limit to first 100 rows
                'analysis': analysis,
                'full_data': df.to_json(),
                'encoding_used': used_encoding
            }
        except Exception as e:
            raise Exception(f"Error processing CSV: {str(e)}")
    
    def _process_excel(self, file_path: str, file_id: str, filename: str) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    sheets_data[sheet_name] = {
                        'data': df.to_dict('records')[:100],
                        'analysis': self._analyze_dataframe(df)
                    }
                except Exception as e:
                    print(f"Error processing sheet {sheet_name}: {e}")
                    sheets_data[sheet_name] = {
                        'data': [],
                        'analysis': {'error': str(e)}
                    }
            
            return {
                'file_id': file_id,
                'filename': filename,
                'type': 'excel',
                'sheets': list(sheets_data.keys()),
                'data': sheets_data,
                'analysis': self._analyze_excel_file(sheets_data)
            }
        except Exception as e:
            raise Exception(f"Error processing Excel: {str(e)}")
    
    def _process_pdf(self, file_path: str, file_id: str, filename: str) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text += page.extract_text() + "\n"
                    except Exception as e:
                        print(f"Error extracting text from page {page_num}: {e}")
            
            return {
                'file_id': file_id,
                'filename': filename,
                'type': 'pdf',
                'text': text,
                'analysis': {
                    'pages': len(pdf_reader.pages),
                    'word_count': len(text.split()),
                    'char_count': len(text)
                }
            }
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _process_docx(self, file_path: str, file_id: str, filename: str) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                'file_id': file_id,
                'filename': filename,
                'type': 'docx',
                'text': text,
                'analysis': {
                    'paragraphs': len(doc.paragraphs),
                    'word_count': len(text.split()),
                    'char_count': len(text)
                }
            }
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def _process_txt(self, file_path: str, file_id: str, filename: str) -> Dict[str, Any]:
        """Process TXT file"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    used_encoding = encoding
                    break
                except:
                    continue
            
            if text is None:
                raise Exception("Could not read text file with any supported encoding")
            
            return {
                'file_id': file_id,
                'filename': filename,
                'type': 'txt',
                'text': text,
                'analysis': {
                    'word_count': len(text.split()),
                    'char_count': len(text),
                    'lines': len(text.splitlines())
                },
                'encoding_used': used_encoding
            }
        except Exception as e:
            raise Exception(f"Error processing TXT: {str(e)}")
    
    def _analyze_dataframe(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Analyze pandas DataFrame and prepare chart data"""
        try:
            # Base analysis information
            analysis: Dict[str, Any] = {
                'shape': df.shape,
                'columns': df.columns.tolist(),
                'dtypes': {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
                'null_counts': df.isnull().sum().to_dict(),
                'summary_stats': {},
                'charts': {}
            }

            # Identify numeric and categorical columns
            numeric_columns = df.select_dtypes(include=[np.number]).columns.tolist()
            categorical_columns = [col for col in df.columns if col not in numeric_columns]

            # Summary statistics for numeric columns
            if numeric_columns:
                stats_df = df[numeric_columns].describe()
                analysis['summary_stats'] = {
                    str(col): {str(stat): val for stat, val in stats_df[col].items()}
                    for col in stats_df.columns
                }

            # Generate histogram data for numeric columns
            for col in numeric_columns:
                try:
                    col_data = df[col].dropna().astype(float)
                    # Freedman–Diaconis rule for bin width
                    q1 = col_data.quantile(0.25)
                    q3 = col_data.quantile(0.75)
                    iqr = q3 - q1
                    bin_width = 2 * iqr / (len(col_data) ** (1/3)) if len(col_data) > 0 else 0
                    if bin_width == 0 or np.isnan(bin_width):
                        bins = 10
                    else:
                        bins = int(np.ceil((col_data.max() - col_data.min()) / bin_width))
                        bins = max(min(bins, 40), 5)  # min 5, max 40 bins
                    counts, bin_edges = np.histogram(col_data, bins=bins)
                    analysis['charts'][col] = {
                        'type': 'histogram',
                        'bins': bin_edges.tolist(),
                        'counts': counts.tolist()
                    }
                except Exception:
                    continue

            # Generate bar chart data for categorical columns
            for col in categorical_columns:
                try:
                    value_counts = df[col].astype(str).value_counts().head(10)
                    if not value_counts.empty:
                        analysis['charts'][col] = {
                            'type': 'bar',
                            'categories': value_counts.index.tolist(),
                            'counts': value_counts.tolist()
                        }
                except Exception:
                    continue

            return analysis
        except Exception as e:
            return {
                'shape': df.shape if df is not None else [0, 0],
                'columns': [],
                'error': str(e)
            }
    
    def _analyze_excel_file(self, sheets_data: Dict) -> Dict[str, Any]:
        """Analyze Excel file with multiple sheets"""
        try:
            total_rows = 0
            total_cols = 0
            
            for sheet_name, sheet_data in sheets_data.items():
                analysis = sheet_data.get('analysis', {})
                if 'shape' in analysis:
                    total_rows += analysis['shape'][0]
                    total_cols += analysis['shape'][1]
            
            return {
                'total_sheets': len(sheets_data),
                'total_rows': total_rows,
                'total_columns': total_cols,
                'sheet_names': list(sheets_data.keys())
            }
        except Exception as e:
            return {
                'total_sheets': len(sheets_data),
                'error': str(e)
            }
